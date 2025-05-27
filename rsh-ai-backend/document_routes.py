import os
import tempfile
from flask import Blueprint, request, jsonify
from werkzeug.utils import secure_filename
import logging
from typing import List, Dict, Any
import numpy as np
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_openai import OpenAIEmbeddings
from pinecone import Pinecone
import uuid
import time

# Configure logging
logger = logging.getLogger(__name__)

# Create blueprint
document_bp = Blueprint('document', __name__)

# Allowed file extensions
ALLOWED_EXTENSIONS = {'txt', 'pdf', 'docx', 'doc'}

def allowed_file(filename):
    return '.' in filename and \
           filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_file(file_path: str) -> str:
    """Extract text from file based on its extension."""
    extension = file_path.rsplit('.', 1)[1].lower()
    
    if extension == 'txt':
        # Simple text file
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    elif extension == 'pdf':
        # PDF file
        try:
            # Coba import dari pypdf (versi baru)
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except ImportError:
                # Fallback ke PyPDF2 (versi lama)
                from PyPDF2 import PdfReader
                reader = PdfReader(file_path)
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
        except Exception as e:
            logger.error(f"Error extracting text from PDF: {str(e)}")
            raise
    
    elif extension in ['docx', 'doc']:
        # Word document
        try:
            try:
                # Coba import dari python-docx
                from docx import Document
                doc = Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                # Fallback ke docx (nama modul lama)
                import docx
                doc = docx.Document(file_path)
                return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"Error extracting text from Word document: {str(e)}")
            raise
    
    else:
        raise ValueError(f"Unsupported file extension: {extension}")

def pad_embedding(embedding, target_dim=1536):
    """Pad or truncate an embedding to the target dimension."""
    if not embedding:
        return [0.0] * target_dim
        
    # Convert to numpy array for easier manipulation
    embedding_array = np.array(embedding)
    current_dim = len(embedding_array)
    
    # If current dimension is already target dimension, return as is
    if current_dim == target_dim:
        return embedding
    
    # If current dimension is larger, truncate
    if current_dim > target_dim:
        return embedding_array[:target_dim].tolist()
    
    # If current dimension is smaller, pad with zeros
    padded = np.zeros(target_dim)
    padded[:current_dim] = embedding_array
    return padded.tolist()

def split_text(text: str) -> List[str]:
    """Split text into chunks."""
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200,
        length_function=len,
    )
    return text_splitter.split_text(text)

def create_embeddings(chunks: List[str], source_name: str) -> List[Dict]:
    """Create embeddings for text chunks."""
    try:
        embeddings_model = OpenAIEmbeddings(
            model="text-embedding-ada-002",
            openai_api_key=os.getenv("OPENAI_API_KEY")
        )
        
        # Process each chunk and ensure it has the correct dimension
        vectors = []
        for i, chunk in enumerate(chunks):
            logger.info(f"Creating embedding for chunk {i+1}/{len(chunks)}")
            embedding = embeddings_model.embed_query(chunk)
            # Ensure embedding has the correct dimension
            padded_embedding = pad_embedding(embedding, 1536)
            
            # Create a unique ID for this chunk
            chunk_id = f"{source_name}_{uuid.uuid4()}"
            
            vectors.append({
                "id": chunk_id,
                "values": padded_embedding,
                "metadata": {
                    "text": chunk,
                    "source": source_name,
                    "chunk_index": i,
                    "timestamp": time.time()
                }
            })
        
        return vectors
    except Exception as e:
        logger.error(f"Error creating embeddings: {str(e)}")
        raise

def upload_to_pinecone(vectors: List[Dict]) -> int:
    """Upload vectors to Pinecone."""
    try:
        # Initialize Pinecone client
        pinecone_api_key = os.getenv("PINECONE_API_KEY")
        if not pinecone_api_key:
            logger.error("PINECONE_API_KEY environment variable not set")
            raise ValueError("Pinecone API key not set")
            
        pc = Pinecone(api_key=pinecone_api_key)
        
        # Get the index
        index_name = os.getenv("PINECONE_INDEX_NAME", "rsh-chatbot-index")
        index = pc.Index(index_name)
        
        # Upsert vectors in batches
        batch_size = 100
        for i in range(0, len(vectors), batch_size):
            batch = vectors[i:i+batch_size]
            logger.info(f"Upserting batch {i//batch_size + 1}/{(len(vectors) + batch_size - 1)//batch_size}")
            index.upsert(vectors=batch)
        
        logger.info(f"Successfully uploaded {len(vectors)} vectors to Pinecone")
        return len(vectors)
    except Exception as e:
        logger.error(f"Error uploading to Pinecone: {str(e)}")
        raise

@document_bp.route('/upload', methods=['POST'])
def upload_document():
    """Upload and index a document."""
    # Check if the post request has the file part
    if 'file' not in request.files:
        return jsonify({"error": "No file part"}), 400
    
    file = request.files['file']
    
    # If user does not select file, browser also
    # submit an empty part without filename
    if file.filename == '':
        return jsonify({"error": "No selected file"}), 400
    
    if file and allowed_file(file.filename):
        try:
            # Create a temporary file to store the uploaded file
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(file.filename)[1]) as temp:
                file.save(temp.name)
                temp_path = temp.name
            
            # Simpan ukuran file sebelum diproses lebih lanjut
            file_size = os.path.getsize(temp_path)
            
            # Extract text from file
            logger.info(f"Extracting text from {file.filename}")
            text = extract_text_from_file(temp_path)
            
            # Split text into chunks
            chunks = split_text(text)
            logger.info(f"Document split into {len(chunks)} chunks")
            
            # Create embeddings
            source_name = secure_filename(file.filename)
            vectors = create_embeddings(chunks, source_name)
            logger.info(f"Created {len(vectors)} embeddings")
            
            # Upload to Pinecone
            count = upload_to_pinecone(vectors)
            
            # Clean up the temporary file setelah semua proses selesai
            try:
                os.unlink(temp_path)
                logger.info(f"Temporary file {temp_path} deleted successfully")
            except Exception as e:
                logger.warning(f"Failed to delete temporary file {temp_path}: {str(e)}")
                # Lanjutkan meskipun gagal menghapus file sementara
            
            return jsonify({
                "success": True,
                "message": "Document indexed successfully",
                "file_name": file.filename,
                "file_size": file_size,
                "chunks": len(chunks)
            }), 200
            
        except Exception as e:
            logger.error(f"Error processing document: {str(e)}")
            return jsonify({"success": False, "message": str(e)}), 500
    
    return jsonify({"success": False, "message": "File type not allowed"}), 400

@document_bp.route('/status', methods=['GET'])
def index_status():
    """Get the status of the Pinecone index."""
    try:
        # Initialize Pinecone client
        pinecone_api_key = os.getenv("PINECONE_API_KEY")
        if not pinecone_api_key:
            return jsonify({"error": "Pinecone API key not set"}), 500
            
        pc = Pinecone(api_key=pinecone_api_key)
        
        # Get the index
        index_name = os.getenv("PINECONE_INDEX_NAME", "rsh-chatbot-index")
        index = pc.Index(index_name)
        
        # Get index stats
        stats = index.describe_index_stats()
        
        return jsonify({
            "status": "Ready",
            "index_name": index_name,
            "vector_count": stats.get("total_vector_count", 0),
            "dimension": stats.get("dimension", 0),
            "metric": "cosine",
            "pod_type": "p1",
            "namespaces": list(stats.get("namespaces", {}).keys())
        }), 200
        
    except Exception as e:
        logger.error(f"Error getting index status: {str(e)}")
        return jsonify({
            "status": "error",
            "error": str(e)
        }), 500
